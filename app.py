import os
import json
import random 
import streamlit as st
import google.generativeai as genai
import tempfile
import io
import pypdfium2 as pdfium
import base64
from docx import Document
import pandas as pd
import plotly.express as px
import streamlit.components.v1 as components
from PIL import Image, ImageOps

# ==========================================
# 🛑 核心配置與 API 初始化
# ==========================================
st.set_page_config(page_title="被動元件專利 AI 戰略分析系統", layout="wide")

api_keys = [
    st.secrets.get("GOOGLE_API_KEY_1", st.secrets.get("GOOGLE_API_KEY", "")),
    st.secrets.get("GOOGLE_API_KEY_2", st.secrets.get("GOOGLE_API_KEY", ""))
]
selected_key = random.choice([k for k in api_keys if k])
if selected_key:
    genai.configure(api_key=selected_key)
model = genai.GenerativeModel('gemini-2.5-flash')

# 🌟 初始化系統所有 Session State
if 'ip_report_content' not in st.session_state: st.session_state.ip_report_content = ""
if 'rd_card_data' not in st.session_state: st.session_state.rd_card_data = None
if 'claim_data_t2' not in st.session_state: st.session_state.claim_data_t2 = None
if 'pdf_bytes_main' not in st.session_state: st.session_state.pdf_bytes_main = None
if 'scanned_pages' not in st.session_state: st.session_state.scanned_pages = {}
if 'ai_analysis_result' not in st.session_state: st.session_state.ai_analysis_result = None
if 'rd_database' not in st.session_state: st.session_state.rd_database = [] 
if 'comp_database' not in st.session_state: st.session_state.comp_database = [] 
if 'active_patent_num' not in st.session_state: st.session_state.active_patent_num = ""

SAVE_DIR = "saved_reports"
if not os.path.exists(SAVE_DIR):
    os.makedirs(SAVE_DIR)

# ==========================================
# 🛑 工具函式庫
# ==========================================
def crop_white_margins(img):
    if img.mode != 'RGB':
        img = img.convert('RGB')
    inv = ImageOps.invert(img)
    bbox = inv.getbbox()
    if bbox:
        padded_bbox = (max(0, bbox[0]-20), max(0, bbox[1]-20), min(img.width, bbox[2]+20), min(img.height, bbox[3]+20))
        return img.crop(padded_bbox)
    return img

def check_password():
    if "password_correct" not in st.session_state:
        st.session_state["password_correct"] = False
    if not st.session_state["password_correct"]:
        st.title("🔒 專利戰情室 - 系統登入")
        pwd = st.text_input("請輸入授權密碼", type="password")
        if pwd == st.secrets.get("APP_PASSWORD", ""): 
            st.session_state["password_correct"] = True
            st.rerun()
        elif pwd:
            st.error("密碼錯誤，請重新輸入！")
        return False
    return True

def create_word_doc(text):
    doc = Document()
    doc.add_heading('專利戰略深度分析報告 (IP 嚴謹版)', 0)
    for para in text.split('\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

if not check_password():
    st.stop()

# ==========================================
# 🚧 全局核心上傳與分析區
# ==========================================
st.title("🔋 被動元件專利 AI 戰略分析系統")
st.markdown("---")

with st.container(border=True):
    st.subheader("📥 第一步：設定案件資訊並上傳 PDF (單篇深度分析用)")
    col_input1, col_input2, col_input3 = st.columns([1.5, 2, 1])

    with col_input1:
        applicant_main = st.text_input("申請人 (對手公司)", placeholder="例如：國巨、華新科", key="main_app")
        patent_num_main = st.text_input("專利號", placeholder="例如：I856744", key="main_num")
        st.session_state.active_patent_num = patent_num_main
        if patent_num_main:
            clean_num_m = ''.join(e for e in patent_num_main if e.isalnum())
            google_patents_url = f"https://patents.google.com/patent/TW{clean_num_m}B" if clean_num_m.upper().startswith('I') else f"https://patents.google.com/patent/TW{clean_num_m}U"
            st.markdown(f"👉 [Google Patents 傳送門 **{patent_num_main}**]({google_patents_url})")

    with col_input2:
        status_main = st.selectbox("目前案件狀態", ["請選擇...", "公開", "公告/核准", "核駁", "撤回", "消滅"], key="main_status")
        uploaded_pdf_main = st.file_uploader("上傳專利 PDF 檔", type=["pdf"], key="main_upload")

    with col_input3:
        st.write("")
        st.write("")
        if st.button("🚀 啟動全局 AI 深度解剖", use_container_width=True, type="primary"):
            if status_main == "請選擇..." or uploaded_pdf_main is None or not patent_num_main:
                st.warning("⚠️ 請確認：1.輸入專利號 2.選擇狀態 3.上傳 PDF")
            else:
                pdf_bytes = uploaded_pdf_main.getvalue()
                st.session_state.pdf_bytes_main = pdf_bytes
                st.session_state.scanned_pages = {} 

                safe_app = "".join(c for c in applicant_main if c.isalnum() or c in (' ', '-', '_')).strip()
                folder_name = safe_app if safe_app else "未分類"
                applicant_dir = os.path.join(SAVE_DIR, folder_name)
                if not os.path.exists(applicant_dir): os.makedirs(applicant_dir)

                clean_num = ''.join(e for e in patent_num_main if e.isalnum())
                file_path = os.path.join(applicant_dir, f"{clean_num}_master.json")

                if os.path.exists(file_path):
                    with st.spinner("載入歷史分析數據..."):
                        with open(file_path, "r", encoding="utf-8") as f:
                            saved_data = json.load(f)
                            st.session_state.ip_report_content = saved_data.get("ip_report", "")
                            st.session_state.rd_card_data = saved_data.get("rd_card", None)
                            st.session_state.claim_data_t2 = saved_data.get("vis_data", None)
                        st.success("⚡ 歷史紀錄載入完成！請切換下方頁籤查看。")
                else:
                    with st.spinner("大腦正在進行全局地毯式搜索，約需 30 秒..."):
                        try:
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                                tmp_file.write(pdf_bytes)
                                tmp_file_path = tmp_file.name
                            gemini_file = genai.upload_file(tmp_file_path)
                            
                            is_utility_model = clean_num.upper().startswith('M')
                            
                            if is_utility_model:
                                claim_analysis_prompt = '''
                                【五、 🧱 具體結構特徵拆解 (Structural Features)】：
                                新型專利保護的是具體形狀與構造。請拆解其最核心的「物理結構、形狀與組裝關係」，不要探討抽象概念。
                                【六、 🪤 迴避設計難度評估】：
                                評估要替換掉這個特定結構或材料層疊設計的難度。這個機構設計或封裝結構是達成該功效的「唯一最佳解」嗎？RD 是否容易用其他常見的製程/電極配置/介電材料來繞開？
                                '''
                            else:
                                claim_analysis_prompt = '''
                                【五、 🧩 獨立項全要件拆解 (Claim Chart)】：
                                最廣獨立項（請求項1）拆解，請以 1. 2. 3. 逐行乾淨條列拆解。
                                在最後一行加上「破口（限縮最嚴格之特徵）：精準點出最容易被迴避的限制條件」。
                                【六、 🪤 附屬項隱藏地雷探測】：
                                以數字條列出具備「具體製程參數、層疊結構、相對位置、或材料成分比例限制」的附屬項。
                                '''

                            prompt_master = f'''
                            【⚠️ 語氣與術語強制校準】：你現在是一位資深被動元件專利代理人與研發主管。你具備材料科學、化學工程與電子電機的碩士學歷，或是在被動元件廠具備 3 到 5 年以上的實務研發經驗，懂得材料科學與固態物理、高分子化學、陶瓷製程與封裝工藝、電力電子與安規標準。請使用被動元件與材料研發黑話。
                            我已經提供了一份被動元件相關的專利 PDF 檔案，請仔細閱讀全文。
                            【補充資訊】申請人：{applicant_main} / 目前法律狀態：{status_main} / 專利類型：{"新型專利" if is_utility_model else "發明專利"}

                            【🔴 輸出格式嚴格要求：純 JSON 格式】
                            {{
                              "rd_card": {{
                                "title": "用一句話總結這項技術",
                                "problem": "原本的設計或製程有什麼缺點",
                                "solution": "本專利用了什麼特殊結構或配方解決",
                                "application": "MLCC/電阻/電感/車用電子",
                                "risk_check": ["破口限制特徵1", "破口限制特徵2", "破口限制特徵3"],
                                "design_avoid_rd": ["具體製程或機構迴避方向1", "具體製程或機構迴避方向2"]
                              }},
                              "vis_data": {{
                                "claims": ["1. 請求項逐句", "第二句..."],
                                "components": [ {{"id": "10", "name": "介電層"}} ],
                                "spec_texts": ["【00xx】實施方式段落1", "【00xx】實施方式段落2"]
                              }},
                              "ip_report": "【一、 🚦 FTO 風險判定】\\n(🔴/🟡/🟢 判定與簡述)\\n\\n【二、 📸 技術核心快照】\\n1. 發明目的\\n2. 核心技術\\n3. 宣稱功效\\n\\n【三、 🏢 研發部門精準派發】\\n\\n【四、 🛑 先前技術與妥協分析 (防禦地雷)】\\n\\n{claim_analysis_prompt}\\n\\n【七、 👁️ 侵權可偵測性評估 (如逆向工程、切片分析難度)】\\n\\n【八、 🕵️‍♂️ 實證功效檢驗 (打假雷達)】\\n\\n【九、 🛡️ 高階迴避設計建議 (防範均等論)】\\n\\n【十、 🧬 技術演進與製程整併雷達】\\n\\n【十一、 🏷️ 元件符號圖面提取字典】"
                            }}
                            備註：vis_data 中的 components 必須去尋找專利最後面的「符號簡單說明」，將所有元件提取出來。
                            '''
                            
                            response = model.generate_content([gemini_file, prompt_master])
                            clean_text = response.text.replace('```json', '').replace('```', '').strip()
                            clean_text = clean_text[clean_text.find('{'):clean_text.rfind('}')+1]
                            master_json = json.loads(clean_text)

                            st.session_state.rd_card_data = master_json.get("rd_card")
                            st.session_state.ip_report_content = master_json.get("ip_report")
                            st.session_state.claim_data_t2 = master_json.get("vis_data")
                            
                            with open(file_path, "w", encoding="utf-8") as f:
                                json.dump({"ip_report": st.session_state.ip_report_content, "rd_card": st.session_state.rd_card_data, "vis_data": st.session_state.claim_data_t2}, f, ensure_ascii=False)
                            st.success("✅ 全局分析完成！請切換下方頁籤查看。")
                            os.remove(tmp_file_path)
                            genai.delete_file(gemini_file.name)
                        except Exception as e:
                            st.error(f"分析失敗：{e}")

st.markdown("<br>", unsafe_allow_html=True)

# 🌟 建立頂層四模式切換
main_tab1, main_tab2, main_tab3, main_tab4 = st.tabs([
    "🧑‍💻 Tab 1 研發：迴避設計大屏 (RD)", 
    "⚖️ Tab 2 智權：法務審查中心 (IP)", 
    "🗺️ Tab 3 戰略：宏觀大數據與快篩 (Excel)", 
    "💡 Tab 4 賦能：研發開源技術庫 (Excel)"
])

# ==========================================
# 🧑‍💻 Tab 1：研發迴避設計大屏 (理想排版實作)
# ==========================================
with main_tab1:
    if not st.session_state.rd_card_data or not st.session_state.pdf_bytes_main:
        st.info("請先於上方上傳 PDF 並啟動分析。")
    else:
        rd_data = st.session_state.rd_card_data
        
        # 🌟 上半部：三卡並列
        col_c1, col_c2, col_c3 = st.columns(3)
        
        with col_c1:
            with st.container(border=True, height=450):
                st.markdown(f"#### 🎯 研發戰略看板")
                st.markdown(f"**{rd_data.get('title', '未知技術')}**")
                
                f_color = "red" if "🔴" in st.session_state.ip_report_content else "orange" if "🟡" in st.session_state.ip_report_content else "green"
                f_text = "🔴 具備威脅" if f_color == "red" else "🟡 需注意" if f_color == "orange" else "🟢 低風險"
                st.markdown(f"**🚦 FTO 燈號：** <span style='color:{f_color}; font-weight:bold; font-size:18px;'>{f_text}</span>", unsafe_allow_html=True)
                st.markdown("---")
                st.markdown(f"**🔥 解決痛點：** {rd_data.get('problem', '')}")
                st.markdown(f"**💡 核心解法：** {rd_data.get('solution', '')}")
                st.markdown(f"**🎯 應用場景：** {rd_data.get('application', '')}")

        with col_c2:
            with st.container(border=True, height=450):
                st.markdown("#### ⚔️ 自家技術 CheckBox 檢核")
                st.caption("請確認我司目前設計是否具備以下「權利要求獨立項特徵」：")
                risk_list = rd_data.get('risk_check', [])
                checked_count = 0
                for i, risk in enumerate(risk_list):
                    if st.checkbox(f"{risk}", key=f"risk_c_{i}"): 
                        checked_count += 1
                
                st.markdown("<br>", unsafe_allow_html=True)
                if len(risk_list) > 0:
                    if checked_count == len(risk_list): 
                        st.markdown("<div style='padding:10px; background-color:#ffebee; color:#c62828; border-radius:5px;'><b>⚠️ 警告：特徵全中，高度侵權風險！</b></div>", unsafe_allow_html=True)
                    elif checked_count > 0: 
                        st.markdown(f"<div style='padding:10px; background-color:#fff8e1; color:#f57f17; border-radius:5px;'><b>注意：命中 {checked_count}/{len(risk_list)} 個特徵，具均等論風險。</b></div>", unsafe_allow_html=True)
                    else: 
                        st.markdown("<div style='padding:10px; background-color:#e8f5e9; color:#2e7d32; border-radius:5px;'><b>🎉 全數未命中，文義迴避成功。</b></div>", unsafe_allow_html=True)

        with col_c3:
            with st.container(border=True, height=450):
                st.markdown("#### 🛡️ 迴避設計建議方向")
                st.caption("針對前述之限制特徵，建議研發之修改方向：")
                for avoid in rd_data.get('design_avoid_rd', []):
                    st.markdown(f"✅ {avoid}")

        st.markdown("---")
        
        # 🌟 下半部：終極滿版雙向連動大屏
        st.markdown("### 🎯 終極雙向連動大屏")
        
        pdf_doc_v = pdfium.PdfDocument(st.session_state.pdf_bytes_main)
        total_pages_v = len(pdf_doc_v)

        col_page, col_btn, _ = st.columns([1, 1, 3])
        with col_page:
            target_page = st.number_input(f"📄 跳至專利圖紙頁碼 (共 {total_pages_v} 頁)", min_value=1, max_value=total_pages_v, value=min(2, total_pages_v), key="vis_page_rd")
        
        page = pdf_doc_v[target_page - 1]
        raw_pil_img = page.render(scale=2.0).to_pil()
        cropped_img = crop_white_margins(raw_pil_img) 
        
        img_byte_arr = io.BytesIO()
        cropped_img.save(img_byte_arr, format='JPEG')
        encoded_img = base64.b64encode(img_byte_arr.getvalue()).decode()
        img_uri = f"data:image/jpeg;base64,{encoded_img}"

        is_scanned = str(target_page) in st.session_state.scanned_pages
        with col_btn:
            st.write("")
            if not is_scanned:
                if st.button(f"🔍 啟動圖片標號鎖定", use_container_width=True, key="btn_scan_rd"):
                    with st.spinner("Gemini Vision 正在鎖定座標..."):
                        try:
                            comp_dict_list = st.session_state.claim_data_t2.get("components", [])
                            known_comps_str = json.dumps(comp_dict_list, ensure_ascii=False)
                            
                            prompt_vision = f'''
                            這是一張專利圖。已知元件表：{known_comps_str}。
                            請找出圖片上「所有肉眼可見的數字標號」，並估算其相對座標(0.0~1.0)。
                            【重要防呆】：如果該頁「沒有任何圖形標號」或是「純文字頁」，請絕對只輸出空的陣列：{{ "hotspots": [] }}。
                            不要有任何多餘的解釋文字。嚴格輸出 JSON 格式。
                            範例：{{ "hotspots": [ {{"number": "31", "name": "端電極", "x_rel": 0.45, "y_rel": 0.55}} ] }}
                            '''
                            
                            response_vis = model.generate_content([cropped_img, prompt_vision])
                            
                            if not response_vis.text:
                                ai_visual_data = []
                            else:
                                clean_text_vis = response_vis.text.replace('```json', '').replace('```', '').strip()
                                start_idx = clean_text_vis.find('{')
                                end_idx = clean_text_vis.rfind('}')
                                if start_idx != -1 and end_idx != -1:
                                    clean_text_vis = clean_text_vis[start_idx:end_idx+1]
                                    ai_visual_data = json.loads(clean_text_vis).get("hotspots", [])
                                else:
                                    ai_visual_data = [] 
                                    
                            st.session_state.scanned_pages[str(target_page)] = ai_visual_data
                            st.rerun()
                        except Exception as e: 
                            st.error(f"視覺解析失敗，可能是該頁面缺乏清晰圖形或 API 回應異常。詳細錯誤：{e}")
            else:
                if not st.session_state.scanned_pages[str(target_page)]:
                    st.warning("⚡ 掃描完成，但此頁面未偵測到任何圖形標號。")
                else:
                    st.success("⚡ 座標已鎖定！請體驗下方雙向連動。")

        # 滿版 HTML 渲染
        if is_scanned:
            ai_visual_data = st.session_state.scanned_pages[str(target_page)]
            comp_dict_list = st.session_state.claim_data_t2.get("components", [])
            claim_lines = st.session_state.claim_data_t2.get("claims", [])
            claim_text_full = "<br><br>".join(claim_lines)
            
            for comp in comp_dict_list:
                c_num = comp.get("id", "")
                c_name = comp.get("name", "")
                replacement = f'<span class="comp-text comp-{c_num}" onmouseover="hoverText(\'{c_num}\')" onmouseout="leaveText(\'{c_num}\')">{c_name} ({c_num})</span>'
                claim_text_full = claim_text_full.replace(f"{c_name} ({c_num})", replacement).replace(c_name, replacement)

            hotspots_html = ""
            for spot in ai_visual_data:
                if spot['name'] != "未知":
                    hotspots_html += f"""
                    <div class="hotspot hotspot-marker-{spot['number']}" id="hotspot-{spot['number']}"
                         style="left: {spot['x_rel']*100}%; top: {spot['y_rel']*100}%;"
                         onmouseover="hoverImage('{spot['number']}', '{spot['name']}')" 
                         onmouseout="leaveImage('{spot['number']}')">
                    </div>
                    """

            html_skeleton = f"""
            <!DOCTYPE html>
            <html>
            <head>
            <style>
                body {{ margin: 0; font-family: sans-serif; background: #fff; }}
                .main-container {{ display: flex; height: 800px; width: 100%; border: 1px solid #ddd; border-radius: 8px; overflow: hidden; }}
                .img-section {{ flex: 6; position: relative; overflow: auto; background: #f8f9fa; border-right: 2px solid #ddd; display: flex; justify-content: center; align-items: flex-start; padding: 10px; }}
                .img-wrapper {{ position: relative; display: inline-block; }}
                .patent-img {{ max-width: 100%; height: auto; display: block; }}
                .hotspot {{ position: absolute; width: 35px; height: 35px; transform: translate(-50%, -50%); border-radius: 50%; cursor: pointer; transition: 0.2s; border: 2px solid transparent; z-index: 10; }}
                .hotspot:hover {{ background: rgba(255, 0, 0, 0.3); border: 2px solid red; box-shadow: 0 0 10px rgba(255,0,0,0.5); z-index: 50; }}
                .hotspot-active {{ background: rgba(255, 255, 0, 0.6) !important; border: 3px solid red !important; box-shadow: 0 0 20px red !important; transform: translate(-50%, -50%) scale(1.3); z-index: 50; }}
                #tooltip {{ display: none; position: absolute; background: rgba(0, 0, 0, 0.8); color: white; padding: 6px 12px; border-radius: 4px; font-size: 14px; z-index: 100; pointer-events: none; white-space: nowrap; }}
                .text-section {{ flex: 4; padding: 20px; overflow-y: auto; font-size: 16px; line-height: 1.8; color: #333; }}
                .comp-text {{ color: #0284c7; font-weight: bold; cursor: pointer; border-bottom: 1px dashed #0284c7; padding: 0 2px; transition: 0.2s; }}
                .highlight-active {{ background-color: #fef08a; color: #b91c1c; border-bottom: none; border-radius: 3px; padding: 2px 4px; }}
            </style>
            </head>
            <body>
            <div class="main-container">
                <div class="img-section" id="img-container">
                    <div class="img-wrapper">
                        <img src="{img_uri}" class="patent-img">
                        {hotspots_html}
                    </div>
                    <div id="tooltip"></div>
                </div>
                <div class="text-section">
                    <div style="font-size:18px; font-weight:bold; color:#1e3a8a; border-bottom:2px solid #eee; padding-bottom:8px; margin-bottom:15px; position:sticky; top:0; background:white; z-index:10;">
                        📜 獨立項文義對應 (雙向連動)
                    </div>
                    {claim_text_full}
                </div>
            </div>
            <script>
                const tooltip = document.getElementById('tooltip');
                function hoverImage(num, name) {{
                    document.onmousemove = e => {{ tooltip.style.left = (e.pageX + 15) + 'px'; tooltip.style.top = (e.pageY + 15) + 'px'; }};
                    tooltip.innerHTML = "標號 <b>" + num + "</b> : " + name; tooltip.style.display = 'block';
                    document.querySelectorAll('.comp-' + num).forEach((el, i) => {{ el.classList.add('highlight-active'); if(i===0) el.scrollIntoView({{behavior:'smooth', block:'center'}}); }});
                }}
                function leaveImage(num) {{
                    document.onmousemove = null; tooltip.style.display = 'none';
                    document.querySelectorAll('.comp-' + num).forEach(el => el.classList.remove('highlight-active'));
                }}
                function hoverText(num) {{
                    document.querySelectorAll('.comp-' + num).forEach(el => el.classList.add('highlight-active'));
                    const hs = document.getElementById('hotspot-' + num);
                    if(hs) {{ hs.classList.add('hotspot-active'); hs.scrollIntoView({{behavior:'smooth', block:'center'}}); }}
                }}
                function leaveText(num) {{
                    document.querySelectorAll('.comp-' + num).forEach(el => el.classList.remove('highlight-active'));
                    const hs = document.getElementById('hotspot-' + num);
                    if(hs) hs.classList.remove('hotspot-active');
                }}
            </script>
            </body>
            </html>
            """
            components.html(html_skeleton, height=820, scrolling=False)

# ==========================================
# ⚖️ Tab 2：智權法務審查中心 (IP)
# ==========================================
with main_tab2:
    if not st.session_state.ip_report_content or not st.session_state.claim_data_t2:
        st.info("請先於上方上傳 PDF 並啟動分析。")
    else:
        st.markdown("## 🏛️ 智權法務審查工作站")
        ip_tab_report, ip_tab_claim = st.tabs(["📄 智權戰略深度報告 (含破口與地雷)", "⚖️ 請求項文義比對 (三視窗)"])
        
        with ip_tab_report:
            col_r1, col_r2 = st.columns([3, 1])
            with col_r1: st.markdown("以下為嚴格遵守指令生成的 11 項實務拆解報告：")
            with col_r2:
                st.download_button(
                    label="📥 下載 Word 報告",
                    data=create_word_doc(st.session_state.ip_report_content),
                    file_name=f"IP_Report_{st.session_state.active_patent_num}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            with st.container(height=650, border=True):
                st.markdown(st.session_state.ip_report_content)
        
        with ip_tab_claim:
            components_list = st.session_state.claim_data_t2.get("components", [])
            if components_list:
                comp_options = {f"[{c.get('id','')}] {c.get('name','')}": c for c in components_list}
                col_sel, _ = st.columns([1, 1])
                with col_sel:
                    selected_comp = st.selectbox(f"🎯 選擇比對目標元件 (共 {len(components_list)} 個)：", list(comp_options.keys()), key="ip_comp_sel")
                    active_c = comp_options[selected_comp]
                
                st.markdown("<br>", unsafe_allow_html=True)
                col_i1, col_i2, col_i3 = st.columns([1.2, 1, 1.2])
                
                with col_i1:
                    st.markdown("### 🖼️ 專利圖面")
                    pdf_doc_ip = pdfium.PdfDocument(st.session_state.pdf_bytes_main)
                    pg_ip = st.number_input("頁碼", min_value=1, max_value=len(pdf_doc_ip), value=min(2, len(pdf_doc_ip)), key="ip_pg")
                    with st.container(height=600, border=True):
                        st.image(pdf_doc_ip[pg_ip - 1].render(scale=2.0).to_pil(), use_container_width=True)
                
                with col_i2:
                    st.markdown("### 🧩 獨立項文義")
                    with st.container(height=600, border=True):
                        for line in st.session_state.claim_data_t2.get("claims", []):
                            if active_c['name'] in line:
                                hl_line = line.replace(active_c['name'], f"<span style='background-color:#fff3cd; font-weight:bold; color:#856404; padding:2px 4px; border-radius:3px;'>{active_c['name']}</span>")
                                st.markdown(f"<div style='padding: 8px; border-bottom: 1px dashed #eee;'>{hl_line}</div>", unsafe_allow_html=True)
                            else:
                                st.markdown(f"<div style='padding: 8px; border-bottom: 1px dashed #eee; color: #555;'>{line}</div>", unsafe_allow_html=True)

                with col_i3:
                    st.markdown("### 📖 說明書具體限制")
                    with st.container(height=600, border=True):
                        st.info(f"📍 目標：**{active_c['name']} ({active_c.get('id','')})**")
                        found_texts = [t for t in st.session_state.claim_data_t2.get('spec_texts', []) if active_c['name'] in t or active_c.get('id','') in t]
                        if not found_texts: st.warning("未找到說明。")
                        else:
                            for t in found_texts:
                                hl_t = t.replace(active_c['name'], f"<mark style='background-color:#cce5ff; color:#004085; font-weight:bold; padding:2px; border-radius:3px;'>{active_c['name']}</mark>")
                                st.markdown(f"<div style='background: #f8f9fa; padding: 10px; border-left: 4px solid #007bff; margin-bottom: 10px;'>{hl_t}</div>", unsafe_allow_html=True)
